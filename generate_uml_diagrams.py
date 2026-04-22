"""
Simple black-and-white UML diagrams — clean hand-drawn style.
6 diagrams in a single A4 Word document.
"""

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

K = '#000000'   # black
W = '#ffffff'   # white

# ── tiny helpers ──────────────────────────────────────────────

def new_fig(w=9, h=12):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')
    ax.axis('off')
    return fig, ax

def save(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf

def rect(ax, cx, cy, w, h, txt='', fs=8, bold=False):
    ax.add_patch(mpatches.FancyBboxPatch(
        (cx-w/2, cy-h/2), w, h, boxstyle='square,pad=0',
        lw=1, edgecolor=K, facecolor=W))
    if txt:
        ax.text(cx, cy, txt, ha='center', va='center', fontsize=fs,
                fontweight='bold' if bold else 'normal', color=K,
                multialignment='center')

def hrect(ax, cx, cy, w, h, txt='', fs=8):
    """Rounded rectangle for state/activity nodes."""
    ax.add_patch(mpatches.FancyBboxPatch(
        (cx-w/2, cy-h/2), w, h, boxstyle='round,pad=0.05',
        lw=1, edgecolor=K, facecolor=W))
    if txt:
        ax.text(cx, cy, txt, ha='center', va='center', fontsize=fs,
                color=K, multialignment='center')

def arr(ax, x0, y0, x1, y1, lbl='', fs=7.5, dashed=False, lw=1.0, dx=0, dy=0.12):
    ls = (0, (4, 3)) if dashed else '-'
    ax.annotate('', xy=(x1, y1), xytext=(x0, y0),
                arrowprops=dict(arrowstyle='->', color=K, lw=lw, linestyle=ls))
    if lbl:
        mx, my = (x0+x1)/2 + dx, (y0+y1)/2 + dy
        ax.text(mx, my, lbl, ha='center', va='bottom', fontsize=fs, color=K,
                multialignment='center')

def vdash(ax, x, y0, y1):
    ax.plot([x, x], [y0, y1], color=K, lw=0.6, ls=(0, (4, 3)))

def hbar(ax, x0, x1, y, lw=3):
    ax.plot([x0, x1], [y, y], color=K, lw=lw)

def dot(ax, cx, cy, r=0.18):
    ax.add_patch(plt.Circle((cx, cy), r, color=K, zorder=4))

def end_dot(ax, cx, cy, r=0.18):
    ax.add_patch(plt.Circle((cx, cy), r, color=K, zorder=4))
    ax.add_patch(plt.Circle((cx, cy), r+0.1, color=K, fill=False, lw=1.5, zorder=4))

def diamond(ax, cx, cy, w=1.5, h=0.55):
    xs = [cx, cx+w/2, cx, cx-w/2, cx]
    ys = [cy+h/2, cy, cy-h/2, cy, cy+h/2]
    ax.plot(xs, ys, color=K, lw=1)
    ax.fill(xs, ys, color=W)

def actor(ax, x, top):
    r = 0.18
    ax.add_patch(plt.Circle((x, top), r, color=K, fill=False, lw=1))
    ax.plot([x, x], [top-r, top-r-0.38], color=K, lw=1)
    ax.plot([x-0.22, x+0.22], [top-r-0.15, top-r-0.15], color=K, lw=1)
    ax.plot([x, x-0.18], [top-r-0.38, top-r-0.65], color=K, lw=1)
    ax.plot([x, x+0.18], [top-r-0.38, top-r-0.65], color=K, lw=1)

def actbar(ax, x, ytop, ybot, w=0.1):
    ax.add_patch(mpatches.FancyBboxPatch(
        (x-w/2, ybot), w, ytop-ybot, boxstyle='square,pad=0',
        lw=0.8, edgecolor=K, facecolor='#e0e0e0'))

def lane_box(ax, lx, rx, ytop, ybot, name, fs=9):
    ax.add_patch(mpatches.FancyBboxPatch(
        (lx, ybot), rx-lx, ytop-ybot, boxstyle='square,pad=0',
        lw=1, edgecolor=K, facecolor=W))
    # header strip
    ax.add_patch(mpatches.FancyBboxPatch(
        (lx, ytop-0.6), rx-lx, 0.6, boxstyle='square,pad=0',
        lw=1, edgecolor=K, facecolor='#dddddd'))
    ax.text((lx+rx)/2, ytop-0.3, name, ha='center', va='center',
            fontsize=fs, fontweight='bold', color=K)

# ═══════════════════════════════════════════════════════════════
# DIAGRAM 1 — Sequence: Recommendation
# ═══════════════════════════════════════════════════════════════
def d1_rec_seq():
    fig, ax = new_fig(10, 13)
    ax.set_xlim(0, 10); ax.set_ylim(0, 13)
    fig.suptitle('Figure 1.1  Sequence Diagram — Movie Recommendation System',
                 fontsize=10, fontweight='bold', y=0.99)

    cols = [1.0, 2.6, 4.5, 6.8, 9.0]
    names = ['User', 'Browser', 'Django View', 'SVD Engine', 'Database']
    top = 12.2

    for x, n in zip(cols, names):
        if n == 'User':
            actor(ax, x, top+0.3)
            ax.text(x, top-0.55, n, ha='center', fontsize=8, fontweight='bold', color=K)
        else:
            rect(ax, x, top, 1.2, 0.5, n, fs=8, bold=True)
        vdash(ax, x, top-0.25, 0.3)

    # activation bars
    actbar(ax, cols[2], 11.5, 2.0)
    actbar(ax, cols[3], 9.2,  5.8)
    actbar(ax, cols[4], 10.5, 9.8)
    actbar(ax, cols[4], 8.5,  7.8)

    msgs = [
        (11.5, cols[0], cols[1], '1: GET /recommendations/'),
        (10.8, cols[1], cols[2], '2: HTTP Request'),
        (10.1, cols[2], cols[4], '3: fetch user ratings'),
        (9.5,  cols[4], cols[2], '4: ratings queryset', True),
        (9.2,  cols[2], cols[3], '5: get_recommendations(user_id)'),
    ]
    for row in msgs:
        y, x0, x1, lbl = row[0], row[1], row[2], row[3]
        dash = len(row) > 4
        arr(ax, x0, y, x1, y, lbl, dashed=dash)

    # loop box
    ax.add_patch(mpatches.FancyBboxPatch((5.6, 6.8), 4.0, 2.0,
        boxstyle='square,pad=0', lw=1, edgecolor=K, facecolor=W, linestyle='--'))
    ax.text(5.65, 8.75, 'loop  [for each candidate movie]', fontsize=7, color=K)

    arr(ax, cols[3], 8.5, cols[4], 8.5, '6: query movie features')
    arr(ax, cols[4], 7.9, cols[3], 7.9, '7: movie data', dashed=True)
    ax.text(cols[3], 7.3, '8: compute SVD score', ha='center', fontsize=7, color=K, style='italic')

    posts = [
        (5.8, cols[3], cols[2], '9: ranked movie list',    True),
        (5.1, cols[2], cols[4], '10: prefetch genres'),
        (4.4, cols[4], cols[2], '11: genre data',          True),
        (3.7, cols[2], cols[1], '12: render template',     True),
        (3.0, cols[1], cols[0], '13: HTML response',       True),
    ]
    for row in posts:
        y, x0, x1, lbl = row[0], row[1], row[2], row[3]
        arr(ax, x0, y, x1, y, lbl, dashed=True)

    ax.text(4.5, 1.6, '[every 50th rating → background thread retrains SVD]',
            ha='center', fontsize=7.5, color=K, style='italic',
            bbox=dict(boxstyle='round,pad=0.3', facecolor=W, edgecolor=K, lw=0.8))

    return save(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 2 — State: Recommendation
# ═══════════════════════════════════════════════════════════════
def d2_rec_state():
    fig, ax = new_fig(8, 13)
    ax.set_xlim(0, 8); ax.set_ylim(0, 13)
    fig.suptitle('Figure 1.2  State Diagram — Movie Recommendation System',
                 fontsize=10, fontweight='bold', y=0.99)

    cx = 4.0
    states = [
        (cx, 12.2, 'START'),
        (cx, 11.0, 'User Visits\n/recommendations'),
        (cx, 9.7,  'Checking User History'),
        (2.2, 8.3, 'Cold Start\n(No History)'),
        (5.8, 8.3, 'Has Ratings'),
        (5.8, 7.0, 'Loading SVD Model'),
        (5.8, 5.7, 'Generating Candidates'),
        (5.8, 4.4, 'Scoring & Ranking'),
        (2.2, 7.0, 'Fetch Trending\n& Top Rated'),
        (cx, 3.1,  'Displaying Recommendations'),
        (cx, 1.8,  'END'),
    ]

    def gxy(lbl):
        for (x, y, l) in states:
            if l == lbl: return x, y
        return 0, 0

    for (x, y, lbl) in states:
        if lbl == 'START':
            dot(ax, x, y)
        elif lbl == 'END':
            end_dot(ax, x, y)
        else:
            hrect(ax, x, y, 2.6, 0.65, lbl, fs=8)

    flows = [
        ('START',                    'User Visits\n/recommendations', ''),
        ('User Visits\n/recommendations', 'Checking User History',   'load session'),
    ]
    for (a, b, lbl) in flows:
        x0,y0 = gxy(a); x1,y1 = gxy(b)
        arr(ax, x0, y0-0.18, x1, y1+0.33, lbl)

    # Decision diamond
    dx, dy = gxy('Checking User History')
    diamond(ax, dx, dy-0.9, w=2.0, h=0.6)
    ax.text(dx, dy-0.9, 'ratings?', ha='center', va='center', fontsize=7.5, color=K)

    # No → cold start
    ax.annotate('', xy=(2.2, 8.63), xytext=(dx-1.0, dy-0.9),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(2.8, 8.2, 'No', fontsize=7.5, color=K)

    # Yes → has ratings
    ax.annotate('', xy=(5.8, 8.63), xytext=(dx+1.0, dy-0.9),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(5.2, 8.2, 'Yes', fontsize=7.5, color=K)

    for (a, b, lbl) in [
        ('Cold Start\n(No History)', 'Fetch Trending\n& Top Rated', ''),
        ('Has Ratings',              'Loading SVD Model',           ''),
        ('Loading SVD Model',        'Generating Candidates',       ''),
        ('Generating Candidates',    'Scoring & Ranking',           ''),
    ]:
        x0,y0 = gxy(a); x1,y1 = gxy(b)
        arr(ax, x0, y0-0.33, x1, y1+0.33, lbl)

    # both → display
    x0,y0 = gxy('Scoring & Ranking')
    x1,y1 = gxy('Displaying Recommendations')
    arr(ax, x0, y0-0.33, x1, y1+0.33, '')
    x0,y0 = gxy('Fetch Trending\n& Top Rated')
    arr(ax, x0, y0-0.33, x1, y1+0.33, '')

    x0,y0 = gxy('Displaying Recommendations')
    x1,y1 = gxy('END')
    arr(ax, x0, y0-0.33, x1, y1+0.28, '')

    # trainingError back-loop on SVD
    lx, ly = gxy('Loading SVD Model')
    ax.annotate('', xy=(lx, ly+0.33), xytext=(7.2, ly+0.33),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.plot([7.2, 7.2], [ly-0.33, ly+0.33], color=K, lw=1)
    ax.annotate('', xy=(7.2, ly-0.33), xytext=(lx+1.3, ly-0.33),
                arrowprops=dict(arrowstyle='-', color=K, lw=1))
    ax.text(7.3, ly, 'trainingError()\n[retry]', fontsize=7, color=K, ha='left',
            va='center', style='italic')

    return save(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 3 — Activity: Recommendation  (3 swimlanes)
# ═══════════════════════════════════════════════════════════════
def d3_rec_activity():
    fig, ax = new_fig(12, 15)
    ax.set_xlim(0, 12); ax.set_ylim(0, 15)
    fig.suptitle('Figure 1.3  Activity Diagram — Movie Recommendation System',
                 fontsize=10, fontweight='bold', y=0.99)

    lanes = [(0, 4, 'User'), (4, 8, 'Server (Django)'), (8, 12, 'SVD Model')]
    for lx, rx, nm in lanes:
        lane_box(ax, lx, rx, 14.5, 0.2, nm)

    def lc(i): return (lanes[i][0]+lanes[i][1])/2

    acts = [
        (0, 13.5, 'Open /recommendations'),
        (1, 12.5, 'Receive HTTP Request'),
        (1, 11.5, 'Query User Rating History'),
        (0,  9.5, 'Show Trending & Top Rated'),
        (2, 10.3, 'Load SVD Model (cached)'),
        (2,  9.3, 'Compute Predicted Ratings'),
        (2,  8.3, 'Sort & Group by Genre'),
        (1,  7.2, 'Fetch Movie Metadata\n(prefetch genres)'),
        (1,  6.0, 'Build Template Context'),
        (0,  4.9, 'Render Recommendations'),
        (0,  3.7, 'User Rates a Movie'),
        (1,  3.7, 'Save Rating to DB'),
        (2,  3.7, 'Retrain SVD\n(background thread)'),
    ]

    ap = {}
    for lane, cy, lbl in acts:
        cx = lc(lane)
        ap[lbl] = (cx, cy)
        hrect(ax, cx, cy, 2.8, 0.6, lbl, fs=7.5)

    dot(ax, lc(0), 14.2)
    arr(ax, lc(0), 14.02, lc(0), 13.8)

    # decision: has ratings?
    dec_cx, dec_cy = lc(1), 10.6
    diamond(ax, dec_cx, dec_cy, w=2.4, h=0.7)
    ax.text(dec_cx, dec_cy, 'Has ratings?', ha='center', va='center', fontsize=7.5, color=K)

    arr(ax, lc(0), 13.2, lc(1), 12.8, '')
    arr(ax, lc(1), 12.2, lc(1), 11.8, '')
    arr(ax, lc(1), 11.2, dec_cx, dec_cy+0.35, '')

    # No → cold start
    ax.annotate('', xy=(lc(0), 9.8), xytext=(dec_cx-1.2, dec_cy),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(2.8, 10.3, 'No', fontsize=7.5, color=K)

    # Yes → SVD
    ax.annotate('', xy=(lc(2), 10.6), xytext=(dec_cx+1.2, dec_cy),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(9.2, 10.3, 'Yes', fontsize=7.5, color=K)

    arr(ax, lc(2), 10.0, lc(2), 9.6, '')
    arr(ax, lc(2), 9.0, lc(2), 8.6, '')

    # merge — both → fetch metadata
    arr(ax, lc(2), 8.0, lc(1), 7.5, '', dy=0.1)
    arr(ax, lc(0), 9.2, lc(1), 7.5, '', dy=0.1)

    arr(ax, lc(1), 6.9, lc(1), 6.3, '')
    arr(ax, lc(1), 5.7, lc(0), 5.2, '')
    arr(ax, lc(0), 4.6, lc(0), 4.0, '')
    arr(ax, lc(0), 3.4, lc(1), 3.4, '')
    arr(ax, lc(1), 3.4, lc(2), 3.4, '')

    end_cx = lc(1)
    end_cy = 2.2
    end_dot(ax, end_cx, end_cy)
    arr(ax, lc(2), 3.1, end_cx, end_cy+0.28, '')

    return save(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 4 — Sequence: Sentiment
# ═══════════════════════════════════════════════════════════════
def d4_sent_seq():
    fig, ax = new_fig(10, 13)
    ax.set_xlim(0, 10); ax.set_ylim(0, 13)
    fig.suptitle('Figure 2.1  Sequence Diagram — Sentiment Analysis System',
                 fontsize=10, fontweight='bold', y=0.99)

    cols = [1.0, 2.6, 4.5, 7.0, 9.2]
    names = ['User', 'Browser', 'Review View', 'Sentiment\nModel', 'Database']
    top = 12.2

    for x, n in zip(cols, names):
        if n == 'User':
            actor(ax, x, top+0.3)
            ax.text(x, top-0.55, n, ha='center', fontsize=8, fontweight='bold', color=K)
        else:
            rect(ax, x, top, 1.3, 0.5, n, fs=8, bold=True)
        vdash(ax, x, top-0.25, 0.3)

    actbar(ax, cols[2], 11.5, 2.5)
    actbar(ax, cols[3], 9.5,  6.0)
    actbar(ax, cols[4], 5.0,  4.3)

    arr(ax, cols[0], 11.5, cols[1], 11.5, '1: write review text')
    arr(ax, cols[1], 10.8, cols[2], 10.8, '2: POST /reviews/add/<id>/')
    ax.text(cols[2], 10.2, '3: validate form', ha='center', fontsize=7.5, color=K, style='italic')

    # alt fragment
    ax.add_patch(mpatches.FancyBboxPatch((5.5, 7.2), 4.2, 2.5,
        boxstyle='square,pad=0', lw=1, edgecolor=K, facecolor=W, linestyle='--'))
    ax.text(5.55, 9.65, 'alt', fontsize=7.5, fontweight='bold', color=K)
    ax.text(5.55, 9.45, '[model cached]', fontsize=7, color=K, style='italic')
    ax.plot([5.5, 9.7], [8.5, 8.5], color=K, lw=0.7, ls='--')
    ax.text(5.55, 8.4, '[load from disk]', fontsize=7, color=K, style='italic')

    arr(ax, cols[2], 9.5, cols[3], 9.5,  '4: analyze(review_text)')
    arr(ax, cols[3], 8.9, cols[2], 8.9,  '5: sentiment + confidence', dashed=True)
    arr(ax, cols[2], 8.0, cols[3], 8.0,  "4': load_sentiment_model()")
    arr(ax, cols[3], 7.4, cols[2], 7.4,  "5': model instance", dashed=True)

    arr(ax, cols[2], 6.5, cols[4], 6.5,  '6: save review + sentiment')
    arr(ax, cols[4], 5.8, cols[2], 5.8,  '7: saved OK', dashed=True)
    arr(ax, cols[2], 5.1, cols[1], 5.1,  '8: JSON {message, sentiment}', dashed=True)
    arr(ax, cols[1], 4.4, cols[0], 4.4,  '9: show toast + update UI', dashed=True)

    ax.text(4.5, 2.0, '[every 50th review → background thread retrains Naive Bayes]',
            ha='center', fontsize=7.5, color=K, style='italic',
            bbox=dict(boxstyle='round,pad=0.3', facecolor=W, edgecolor=K, lw=0.8))

    return save(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 5 — State: Sentiment
# ═══════════════════════════════════════════════════════════════
def d5_sent_state():
    fig, ax = new_fig(8, 13)
    ax.set_xlim(0, 8); ax.set_ylim(0, 13)
    fig.suptitle('Figure 2.2  State Diagram — Sentiment Analysis System',
                 fontsize=10, fontweight='bold', y=0.99)

    cx = 4.0
    states = [
        (cx,  12.3, 'START'),
        (cx,  11.2, 'Review Form Displayed'),
        (cx,  10.0, 'User Typing Review'),
        (cx,   8.8, 'Form Submitted'),
        (cx,   7.6, 'Validating Input'),
        (2.0,  6.4, 'Validation\nFailed'),
        (cx,   6.4, 'Loading Sentiment\nModel'),
        (cx,   5.2, 'Analysing Sentiment'),
        (cx,   4.0, 'Saving Review\nto Database'),
        (2.5,  2.8, 'Positive\nReview Saved'),
        (5.5,  2.8, 'Negative\nReview Saved'),
        (cx,   1.6, 'Response Returned'),
        (cx,   0.4, 'END'),
    ]

    def gxy(lbl):
        for x,y,l in states:
            if l == lbl: return x,y
        return 0,0

    for x,y,lbl in states:
        if lbl == 'START': dot(ax, x, y)
        elif lbl == 'END': end_dot(ax, x, y)
        else: hrect(ax, x, y, 2.6, 0.65, lbl, fs=8)

    for a,b in [('START','Review Form Displayed'),
                ('Review Form Displayed','User Typing Review'),
                ('User Typing Review','Form Submitted'),
                ('Form Submitted','Validating Input')]:
        x0,y0 = gxy(a); x1,y1 = gxy(b)
        arr(ax, x0, y0-0.18 if a=='START' else y0-0.33, x1, y1+0.33)

    # decision after Validating
    vx,vy = gxy('Validating Input')
    diamond(ax, vx, vy-0.85, w=2.0, h=0.6)
    ax.text(vx, vy-0.85, 'valid?', ha='center', va='center', fontsize=7.5, color=K)

    # No → failed
    ax.annotate('', xy=(2.0, 6.72), xytext=(vx-1.0, vy-0.85),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(2.7, 6.5, 'No', fontsize=7.5)

    # Yes → loading
    ax.annotate('', xy=(cx, 6.72), xytext=(vx+1.0, vy-0.85),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(4.6, 6.5, 'Yes', fontsize=7.5)

    # invalidText() back-loop
    ax.annotate('', xy=(cx-1.3, 10.32), xytext=(0.8, 6.72),
                arrowprops=dict(arrowstyle='->', color=K, lw=1,
                                connectionstyle='arc3,rad=0.35'))
    ax.text(0.5, 8.6, 'invalidText()\n[re-enter]', fontsize=7, ha='center',
            color=K, style='italic')

    for a,b in [('Loading Sentiment\nModel','Analysing Sentiment'),
                ('Analysing Sentiment','Saving Review\nto Database')]:
        x0,y0 = gxy(a); x1,y1 = gxy(b)
        arr(ax, x0, y0-0.33, x1, y1+0.33)

    sx,sy = gxy('Saving Review\nto Database')
    ax.annotate('', xy=(2.5, 3.12), xytext=(sx-0.7, sy-0.33),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(2.8, 3.5, '+ve', fontsize=7.5)
    ax.annotate('', xy=(5.5, 3.12), xytext=(sx+0.7, sy-0.33),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(5.2, 3.5, '-ve', fontsize=7.5)

    rx,ry = gxy('Response Returned')
    ax.annotate('', xy=(rx-0.8, ry+0.33), xytext=(2.5, 2.47),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.annotate('', xy=(rx+0.8, ry+0.33), xytext=(5.5, 2.47),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))

    x0,y0 = gxy('Response Returned')
    arr(ax, x0, y0-0.33, cx, 0.58)

    return save(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 6 — Activity: Sentiment (3 swimlanes)
# ═══════════════════════════════════════════════════════════════
def d6_sent_activity():
    fig, ax = new_fig(12, 15)
    ax.set_xlim(0, 12); ax.set_ylim(0, 15)
    fig.suptitle('Figure 2.3  Activity Diagram — Sentiment Analysis System',
                 fontsize=10, fontweight='bold', y=0.99)

    lanes = [(0, 4, 'User'), (4, 8, 'Web Server (Django)'), (8, 12, 'Analysis Model')]
    for lx, rx, nm in lanes:
        lane_box(ax, lx, rx, 14.5, 0.2, nm)

    def lc(i): return (lanes[i][0]+lanes[i][1])/2

    acts = [
        (0, 13.5, 'Open Movie Detail Page'),
        (0, 12.4, 'Fill in Review Form'),
        (0, 11.3, 'Click Submit Review'),
        (1, 11.3, 'Receive POST Request'),
        (1, 10.2, 'Validate Form Data'),
        (0,  8.3, 'Show Validation Errors'),
        (2,  9.5, 'Tokenise &\nPreprocess Text'),
        (2,  8.3, 'Load Naive Bayes\nModel (cached)'),
        (2,  7.1, 'Compute Class\nProbabilities'),
        (2,  5.9, 'Return Label +\nConfidence Score'),
        (1,  4.9, 'Save Review &\nSentiment to DB'),
        (0,  3.0, 'Display Sentiment\nBadge & Toast'),
        (2,  3.7, 'Retrain Naive Bayes\n(background thread)'),
    ]

    ap = {}
    for lane, cy, lbl in acts:
        cx = lc(lane)
        ap[lbl] = (cx, cy)
        hrect(ax, cx, cy, 2.8, 0.6, lbl, fs=7.5)

    dot(ax, lc(0), 14.2)
    arr(ax, lc(0), 14.02, lc(0), 13.8)
    arr(ax, lc(0), 13.2,  lc(0), 12.7)
    arr(ax, lc(0), 12.1,  lc(0), 11.6)
    arr(ax, lc(0), 11.0,  lc(1), 11.0)
    arr(ax, lc(1), 11.0,  lc(1), 10.5)

    # decision valid?
    dec_cx, dec_cy = lc(1), 9.5
    diamond(ax, dec_cx, dec_cy, w=2.4, h=0.7)
    ax.text(dec_cx, dec_cy, 'Valid?', ha='center', va='center', fontsize=7.5, color=K)
    arr(ax, lc(1), 9.9, dec_cx, dec_cy+0.35)

    # No
    ax.annotate('', xy=(lc(0), 8.6), xytext=(dec_cx-1.2, dec_cy),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(2.5, 9.1, 'No', fontsize=7.5, color=K)

    # re-enter loop
    ax.annotate('', xy=(lc(0), 12.1), xytext=(lc(0), 8.0),
                arrowprops=dict(arrowstyle='->', color=K, lw=1,
                                connectionstyle='arc3,rad=-0.5'))
    ax.text(0.2, 10.0, 're-enter', fontsize=7, ha='left', color=K, style='italic')

    # Yes → model
    ax.annotate('', xy=(lc(2), 9.8), xytext=(dec_cx+1.2, dec_cy),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(9.2, 9.1, 'Yes', fontsize=7.5, color=K)

    arr(ax, lc(2), 9.2, lc(2), 8.6)
    arr(ax, lc(2), 8.0, lc(2), 7.4)
    arr(ax, lc(2), 6.8, lc(2), 6.2)

    # JOIN BAR
    join_y = 5.45
    hbar(ax, lc(2)-0.15, lc(2)+0.15, join_y, lw=3)
    arr(ax, lc(2), 5.6, lc(2), join_y+0.01)
    arr(ax, lc(2), join_y, lc(1), 5.2)

    arr(ax, lc(1), 4.6, lc(1), 4.0)

    # retrain decision
    rd_cx, rd_cy = lc(1), 3.5
    diamond(ax, rd_cx, rd_cy, w=2.6, h=0.7)
    ax.text(rd_cx, rd_cy, 'Every 50th?', ha='center', va='center', fontsize=7.5, color=K)
    arr(ax, lc(1), 4.0, rd_cx, rd_cy+0.35)

    # No → user display
    ax.annotate('', xy=(lc(0), 3.3), xytext=(rd_cx-1.3, rd_cy),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(2.5, 3.15, 'No', fontsize=7.5)

    # Yes → retrain
    ax.annotate('', xy=(lc(2), 4.0), xytext=(rd_cx+1.3, rd_cy),
                arrowprops=dict(arrowstyle='->', color=K, lw=1))
    ax.text(9.3, 3.15, 'Yes', fontsize=7.5)

    # retrain → display
    ax.annotate('', xy=(lc(0), 3.3), xytext=(lc(2), 3.4),
                arrowprops=dict(arrowstyle='->', color=K, lw=1,
                                connectionstyle='arc3,rad=0.25'))

    end_dot(ax, lc(0), 1.8)
    arr(ax, lc(0), 2.7, lc(0), 2.08)

    return save(fig)


# ═══════════════════════════════════════════════════════════════
# BUILD DOCX
# ═══════════════════════════════════════════════════════════════
def build_docx(imgs):
    doc = Document()
    for sec in doc.sections:
        sec.page_width   = Cm(21.0)
        sec.page_height  = Cm(29.7)
        sec.top_margin   = Cm(1.5)
        sec.bottom_margin= Cm(1.5)
        sec.left_margin  = Cm(1.8)
        sec.right_margin = Cm(1.8)

    def heading(txt, lvl=1):
        p = doc.add_heading(txt, level=lvl)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.color.rgb = RGBColor(0,0,0)

    def caption(txt):
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.runs[0]; r.italic = True
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0,0,0)

    # title page
    heading('NextReel — UML Diagrams', 0)
    p = doc.add_paragraph('Movie Recommendation & Sentiment Analysis System')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0,0,0)
    doc.add_page_break()

    sections = [
        ('1  Movie Recommendation System', [
            ('1.1  Sequence Diagram',  'r1', 'Figure 1.1 — Sequence Diagram: Movie Recommendation System'),
            ('1.2  State Diagram',     'r2', 'Figure 1.2 — State Diagram: Movie Recommendation System'),
            ('1.3  Activity Diagram',  'r3', 'Figure 1.3 — Activity Diagram: Movie Recommendation System'),
        ]),
        ('2  Sentiment Analysis System', [
            ('2.1  Sequence Diagram',  's1', 'Figure 2.1 — Sequence Diagram: Sentiment Analysis System'),
            ('2.2  State Diagram',     's2', 'Figure 2.2 — State Diagram: Sentiment Analysis System'),
            ('2.3  Activity Diagram',  's3', 'Figure 2.3 — Activity Diagram: Sentiment Analysis System'),
        ]),
    ]

    first = True
    for sec_title, diagrams in sections:
        if not first: doc.add_page_break()
        heading(sec_title, 1)
        first = False
        for sub_title, key, cap in diagrams:
            doc.add_page_break()
            heading(sub_title, 2)
            doc.add_picture(imgs[key], width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption(cap)

    path = r'e:\Projects\Manoj\NextReel\NextReel_UML_Diagrams_v2.docx'
    doc.save(path)
    return path


if __name__ == '__main__':
    steps = [
        ('r1', '1/6 Sequence — Recommendation', d1_rec_seq),
        ('r2', '2/6 State — Recommendation',    d2_rec_state),
        ('r3', '3/6 Activity — Recommendation',  d3_rec_activity),
        ('s1', '4/6 Sequence — Sentiment',       d4_sent_seq),
        ('s2', '5/6 State — Sentiment',          d5_sent_state),
        ('s3', '6/6 Activity — Sentiment',        d6_sent_activity),
    ]
    imgs = {}
    for key, msg, fn in steps:
        print(msg, '...')
        imgs[key] = fn()

    print('Assembling Word document ...')
    path = build_docx(imgs)
    print('Saved:', path)
